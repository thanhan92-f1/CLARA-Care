from __future__ import annotations

import base64
import io
import re
from typing import Any

import httpx
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

_MERMAID_RENDER_TIMEOUT_SECONDS = 8.0
_MERMAID_MAX_BYTES = 4_500_000
_MERMAID_INFO_PATTERN = re.compile(r"^mermaid(\s+|$)", flags=re.IGNORECASE)


def build_docx_bytes_from_markdown(markdown_text: str) -> bytes:
    doc = Document()
    _configure_document_styles(doc)

    md = MarkdownIt("commonmark", {"html": False}).enable("table")
    tokens = md.parse(_normalize_markdown(markdown_text))
    _render_markdown_tokens(doc, tokens)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _configure_document_styles(doc: DocumentObject) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal_style.font.size = Pt(11)


def _normalize_markdown(value: str) -> str:
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n")


def _render_markdown_tokens(doc: DocumentObject, tokens: list[Token]) -> None:
    index = 0
    while index < len(tokens):
        token = tokens[index]
        token_type = token.type

        if token_type == "heading_open":
            level = _heading_level(token.tag)
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            paragraph = doc.add_heading(level=level)
            _append_inline_tokens(paragraph, inline.children if inline else [])
            index += 3
            continue

        if token_type == "paragraph_open":
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            paragraph = doc.add_paragraph()
            _append_inline_tokens(paragraph, inline.children if inline else [])
            index += 3
            continue

        if token_type in {"bullet_list_open", "ordered_list_open"}:
            ordered = token_type == "ordered_list_open"
            index = _render_list_block(doc, tokens, index, ordered=ordered)
            continue

        if token_type == "fence":
            info = (token.info or "").strip()
            code = token.content or ""
            if _MERMAID_INFO_PATTERN.match(info):
                _append_mermaid_block(doc, code)
            else:
                _append_code_block(doc, code)
            index += 1
            continue

        if token_type == "table_open":
            index = _render_table_block(doc, tokens, index)
            continue

        if token_type == "hr":
            doc.add_paragraph("────────────────────────")
            index += 1
            continue

        if token_type in {"blockquote_open", "blockquote_close", "paragraph_close"}:
            index += 1
            continue

        index += 1


def _heading_level(tag: str) -> int:
    try:
        parsed = int(str(tag or "h2").replace("h", ""))
    except ValueError:
        parsed = 2
    return max(1, min(parsed, 4))


def _render_list_block(
    doc: DocumentObject,
    tokens: list[Token],
    start_index: int,
    *,
    ordered: bool,
) -> int:
    closing_type = "ordered_list_close" if ordered else "bullet_list_close"
    style_name = "List Number" if ordered else "List Bullet"

    index = start_index + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == closing_type:
            return index + 1

        if token.type == "list_item_open":
            item_inline: Token | None = None
            cursor = index + 1
            while cursor < len(tokens) and tokens[cursor].type != "list_item_close":
                current = tokens[cursor]
                if current.type == "inline":
                    item_inline = current
                    break
                cursor += 1
            paragraph = doc.add_paragraph(style=style_name)
            _append_inline_tokens(paragraph, item_inline.children if item_inline else [])
            index = cursor

        index += 1

    return index


def _render_table_block(doc: DocumentObject, tokens: list[Token], start_index: int) -> int:
    rows: list[list[str]] = []
    current_row: list[str] = []
    index = start_index + 1

    while index < len(tokens):
        token = tokens[index]
        token_type = token.type

        if token_type == "table_close":
            break

        if token_type == "tr_open":
            current_row = []
            index += 1
            continue

        if token_type == "tr_close":
            if current_row:
                rows.append(current_row)
            current_row = []
            index += 1
            continue

        if token_type in {"th_open", "td_open"}:
            cell_value = ""
            cursor = index + 1
            while cursor < len(tokens) and tokens[cursor].type not in {"th_close", "td_close"}:
                if tokens[cursor].type == "inline":
                    cell_value = _flatten_inline_text(tokens[cursor].children or [])
                cursor += 1
            current_row.append(cell_value.strip())
            index = cursor + 1
            continue

        index += 1

    if not rows:
        return index + 1

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            table.cell(row_index, column_index).text = value

    return index + 1


def _append_inline_tokens(paragraph: Any, children: list[Token] | None) -> None:
    if not children:
        return

    bold = False
    italic = False
    link_href: str | None = None

    for child in children:
        token_type = child.type

        if token_type == "strong_open":
            bold = True
            continue
        if token_type == "strong_close":
            bold = False
            continue
        if token_type == "em_open":
            italic = True
            continue
        if token_type == "em_close":
            italic = False
            continue
        if token_type == "link_open":
            link_href = _token_attr(child, "href")
            continue
        if token_type == "link_close":
            link_href = None
            continue
        if token_type in {"softbreak", "hardbreak"}:
            paragraph.add_run("\n")
            continue
        if token_type == "code_inline":
            run = paragraph.add_run(child.content or "")
            run.bold = bold
            run.italic = italic
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
            continue
        if token_type != "text":
            continue

        run = paragraph.add_run(child.content or "")
        run.bold = bold
        run.italic = italic
        if link_href:
            run.underline = True
            run.font.color.rgb = RGBColor(0x05, 0x61, 0xC3)


def _flatten_inline_text(children: list[Token]) -> str:
    chunks: list[str] = []
    for child in children:
        if child.type == "text":
            chunks.append(child.content or "")
        elif child.type == "code_inline":
            chunks.append(child.content or "")
        elif child.type in {"softbreak", "hardbreak"}:
            chunks.append(" ")
    return "".join(chunks)


def _append_code_block(doc: DocumentObject, code: str) -> None:
    lines = (code or "").splitlines() or [""]
    for line in lines:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10.5)


def _append_mermaid_block(doc: DocumentObject, code: str) -> None:
    title = doc.add_paragraph("Mermaid Diagram")
    title.runs[0].bold = True

    png_bytes = _render_mermaid_png(code)
    if png_bytes:
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.2))
        return

    warning = doc.add_paragraph("[Mermaid render unavailable - included as code]")
    warning.runs[0].italic = True
    _append_code_block(doc, code)


def _render_mermaid_png(code: str) -> bytes | None:
    payload = (code or "").strip()
    if not payload:
        return None

    encoded = base64.urlsafe_b64encode(payload.encode("utf-8")).decode("ascii").rstrip("=")
    url = f"https://mermaid.ink/img/{encoded}"

    try:
        with httpx.Client(timeout=_MERMAID_RENDER_TIMEOUT_SECONDS, follow_redirects=True) as client:
            response = client.get(url)
            if response.status_code != 200:
                return None
            content_type = str(response.headers.get("content-type") or "").lower()
            if not content_type.startswith("image/"):
                return None
            data = response.content
            if not data or len(data) > _MERMAID_MAX_BYTES:
                return None
            return data
    except Exception:
        return None


def _token_attr(token: Token, key: str) -> str | None:
    attrs = token.attrs or {}
    if isinstance(attrs, dict):
        value = attrs.get(key)
        return str(value) if value is not None else None
    if isinstance(attrs, list):
        for entry in attrs:
            if isinstance(entry, (list, tuple)) and len(entry) == 2 and entry[0] == key:
                return str(entry[1])
    return None
